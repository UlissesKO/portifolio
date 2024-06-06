import docx

#Tem que ajeitar a formatação

class Doc():
    def __init__(self, doc):
        self.doc = docx.Document(doc)
    
    def addFormacao(self, txt):
        #Variavel para contar
        a= 0
        #Vai ler todo o arquivo
        for i in self.doc.paragraphs:
            a += 1
            if i.text == "Cursos":
                a -= 1
                break
        #Cria um novo paragrafo no final do doc
        new_paragraph = self.doc.add_paragraph(txt)

        # Move o novo parágrafo para depois do primeiro parágrafo
        self.doc.element.body.insert(a, new_paragraph._element)

        #Salva o documento
        self.doc.save("Atualizador_Currículo\CurriculoTeste.docx")

    def addCurso(self, txt):
        #Variavel para contar
        a= 0
        #Vai ler todo o arquivo
        for i in self.doc.paragraphs:
            a += 1
            if i.text == "Conhecimentos":
                a -= 1
                break
        #Cria um novo paragrafo no final do doc
        new_paragraph = self.doc.add_paragraph(txt)

        # Move o novo parágrafo para depois do primeiro parágrafo
        self.doc.element.body.insert(a, new_paragraph._element)

        #Salva o documento
        self.doc.save("Atualizador_Currículo\CurriculoTeste.docx")

    def addConhecimento(self):
        #Pegar a ultima linha do doc
        pass

obj = Doc("Atualizador_Currículo\CurriculoTeste.docx")


