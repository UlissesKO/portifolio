import docx

#Seta o doc como o documento desejado
doc = docx.Document("Atualizador_Currículo\CurriculoTeste.docx")

#Vai mostrar o paragrafo desejado
print(doc.paragraphs[9].text)
#Cada \n é um paragrafo diferente
print("_________________________________________")
#Toda vez que muda a formatação vira uma "run" diferente
print(doc.paragraphs[9].runs[0].text)
print("_________________________________________")

a= 0
#Vai ler todo o arquivo
for i in doc.paragraphs:
    print(i.text)
    a += 1
    print("\n")
    if i.text == "Conhecimentos":
        print(a)
        break

